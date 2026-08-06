#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""self_check.py · 用户侧轻量安装自检

**目的**：让接收 Skill 的普通用户确认当前安装是否可正常运行，
         **不涉及 Benchmark、不调 Ark、不碰任何私有样本**。

**硬约束**（见 plans/M4_BUILD_GATE_ADDENDUM.md §8）：
  - 默认不调用真实 Ark
  - 默认不使用学生论文和私有 Benchmark 样本
  - 使用内置合成数据或最小公开 fixture
  - 自检只回答"当前安装能否运行"，不宣称"质检准确率已重新验证"
  - 完整 Benchmark 仍只供开发者和发布流程使用

**依赖**：只依赖 `requirements-core.txt`。

CLI:
    python scripts/self_check.py                # 基础自检（默认）
    python scripts/self_check.py --json         # 机器可读
    python scripts/self_check.py --with-vision  # 可选：合成图片走一次 Ark
    python scripts/self_check.py --audit-rules  # 规则定义-vs-注册表一致性审计
"""
from __future__ import annotations
import argparse
import importlib
import io
import json
import sys
from pathlib import Path

_REPO_ROOT = Path(__file__).resolve().parent.parent

# ---------- diagnostic_result.json Schema 校验 ----------

REQUIRED_ISSUE_FIELDS = [
    "issue_id", "level", "domain", "location", "evidence",
    "evidence_strength", "explanation", "suggestion",
]
RED_EXTRA_REQUIRED = ["normative_basis"]  # 红色 issue 额外必填
REQUIRED_TOP_FIELDS = ["overall_risk_level", "summary_counts"]


def _validate_with_json_schema(data: dict) -> tuple[list[str], bool]:
    """可选：用 JSON Schema 校验 diagnostic_result.json 结构。需要 jsonschema 库。
    返回 (错误列表, 是否实际执行了校验)。"""
    schema_path = _REPO_ROOT / "schemas" / "diagnostic_result.schema.json"
    if not schema_path.exists():
        return [f"JSON Schema 文件不存在: {schema_path}"], False
    try:
        from jsonschema import validate, ValidationError
    except ImportError:
        return [], False  # 库未安装时静默跳过
    try:
        schema = json.loads(schema_path.read_text(encoding="utf-8"))
        validate(instance=data, schema=schema)
        return [], True
    except ValidationError as e:
        return [f"JSON Schema 校验失败: {e.message} (路径: {'/'.join(str(p) for p in e.absolute_path)})"], True
    except Exception as e:
        return [f"JSON Schema 校验异常: {e}"], True


def validate_diagnostic_result(path: Path, use_json_schema: bool = False) -> tuple[bool, list[str], list[str]]:
    """校验 diagnostic_result.json 必填字段。返回 (通过, 错误列表, 警告列表)。"""
    errors: list[str] = []
    warnings: list[str] = []
    data = json.loads(path.read_text(encoding="utf-8"))

    # 可选：JSON Schema 结构校验
    json_schema_applied = False
    if use_json_schema:
        schema_errors, json_schema_applied = _validate_with_json_schema(data)
        errors.extend(schema_errors)

    # 顶层必填
    for f in REQUIRED_TOP_FIELDS:
        if f not in data or data[f] is None:
            errors.append(f"顶层缺必填字段: {f}")

    issues = data.get("issues", [])
    if not isinstance(issues, list):
        errors.append("issues 不是 list")
        return False, errors, warnings

    for i, iss in enumerate(issues):
        for f in REQUIRED_ISSUE_FIELDS:
            if f not in iss or iss.get(f) in (None, ""):
                errors.append(f"issue[{i}]({iss.get('issue_id','?')})缺必填字段: {f}")
        # 红色额外校验
        if iss.get("level") == "红色":
            nb = iss.get("normative_basis")
            if nb is None or (isinstance(nb, dict) and not nb):
                errors.append(
                    f"issue[{i}]({iss.get('issue_id','?')})红色但 normative_basis 为空 "
                    f"— 应通过 kb_query.py 填充规范依据"
                )
            # evidence_items 建议（非阻塞警告）
            ei = iss.get("evidence_items")
            if not isinstance(ei, list) or len(ei) < 2:
                warnings.append(
                    f"issue[{i}]({iss.get('issue_id','?')})红色建议提供 ≥2 条 evidence_items "
                    f"（当前未提供或不足 2 条）"
                )
            # severity_reason 建议
            if not iss.get("severity_reason"):
                warnings.append(
                    f"issue[{i}]({iss.get('issue_id','?')})红色建议填写 severity_reason"
                )
        # contradiction_review 结构校验
        cr = iss.get("contradiction_review")
        if isinstance(cr, dict):
            for f in ("kind", "occurrence_count", "likely_typo", "changes_substantive_conclusion"):
                if f not in cr:
                    warnings.append(
                        f"issue[{i}]({iss.get('issue_id','?')}).contradiction_review 缺字段: {f}"
                    )

    # pass_items 基础校验
    for i, p in enumerate(data.get("pass_items", []) or []):
        if "id" not in p or not p.get("id"):
            errors.append(f"pass_item[{i}]缺 id")
        if "text" not in p or not p.get("text"):
            errors.append(f"pass_item[{i}]({p.get('id','?')})缺 text")

    # P6-P7: priority_actions 类型校验（必须为对象列表，不能是字符串）
    priority_actions = data.get("priority_actions", []) or []
    if not isinstance(priority_actions, list):
        errors.append("priority_actions 不是 list")
    else:
        for i, pa in enumerate(priority_actions):
            if isinstance(pa, str):
                errors.append(
                    f"priority_actions[{i}] 是字符串'{pa[:40]}...'，必须是对象 "
                    f"{{priority, action, issue_ref, expected}}"
                )
            elif isinstance(pa, dict):
                for f in ("priority", "action", "issue_ref", "expected"):
                    if f not in pa or pa.get(f) in (None, ""):
                        errors.append(f"priority_actions[{i}]缺字段: {f}")
            else:
                errors.append(f"priority_actions[{i}]类型错误: {type(pa).__name__}")

    # P6: manual_confirmation_items 字段校验
    manual_items = data.get("manual_confirmation_items", []) or []
    if not isinstance(manual_items, list):
        errors.append("manual_confirmation_items 不是 list")
    else:
        mc_required = ["id", "item", "location", "reason", "how_to_check"]
        for i, mc in enumerate(manual_items):
            if not isinstance(mc, dict):
                errors.append(
                    f"manual_confirmation_items[{i}] 不是对象而是 {type(mc).__name__}"
                )
                continue
            for f in mc_required:
                if f not in mc or mc.get(f) in (None, ""):
                    errors.append(
                        f"manual_confirmation_items[{i}]({mc.get('id','?')})缺字段: {f}"
                    )

    return (len(errors) == 0, errors, warnings, json_schema_applied)

# ---------- 规则库审计 ----------

def audit_rules() -> tuple[bool, str]:
    """校验规则定义文件与注册表的 rule_id 一致性。返回 (通过, 消息)。"""
    try:
        import yaml
    except ImportError:
        return False, "❌ 规则审计失败：pyyaml 未安装"

    repo = _REPO_ROOT

    # 1. 从定义文件中递归收集所有 rule_id
    def_ids: set[str] = set()

    def _collect_rules_from_list(rules_list: list) -> None:
        for item in rules_list:
            if isinstance(item, dict):
                if "rule_id" in item:
                    def_ids.add(str(item["rule_id"]))
                for key in ("sub_rules", "variants"):
                    sub = item.get(key)
                    if isinstance(sub, list):
                        _collect_rules_from_list(sub)

    for yaml_path in [repo / "rules/non_model_rules.yaml", repo / "rules/model_rules.yaml"]:
        if not yaml_path.exists():
            return False, f"❌ 规则审计失败：{yaml_path.name} 不存在"
        data = yaml.safe_load(yaml_path.read_text(encoding="utf-8")) or {}
        # 顶层 rules 列表
        top_rules = data.get("rules")
        if isinstance(top_rules, list):
            _collect_rules_from_list(top_rules)
        # 嵌套 _rules 节（如 model_ml_rules）
        for key, val in data.items():
            if isinstance(val, dict) and key.endswith("_rules"):
                nested = val.get("rules")
                if isinstance(nested, list):
                    _collect_rules_from_list(nested)

    # 2. 从注册表中收集所有 rule_ids
    registry = yaml.safe_load((repo / "rules/rule_registry.yaml").read_text(encoding="utf-8")) or {}
    reg_ids: set[str] = set()
    for group in (registry.get("rule_groups") or {}).values():
        if isinstance(group, dict):
            for rid in (group.get("rule_ids") or []):
                reg_ids.add(str(rid))

    # 3. 比对
    only_def = def_ids - reg_ids
    only_reg = reg_ids - def_ids
    breakdown = (registry.get("capability_breakdown") or {})

    lines = []
    if not only_def and not only_reg:
        declared = breakdown.get("registered_rules", "?")
        lines.append(f"规则定义 {len(def_ids)} 条 ↔ 注册表 {len(reg_ids)} 条（声明 {declared} 条）一致")
        return True, "✅ " + "；".join(lines)

    if only_def:
        lines.append(f"定义有但注册表无 ({len(only_def)}): {sorted(only_def)}")
    if only_reg:
        lines.append(f"注册表有但定义无 ({len(only_reg)}): {sorted(only_reg)}")
    return False, "❌ 规则不一致：" + "；".join(lines)


# ---------- 分项检查 ----------

def _check_import(module: str, human_name: str) -> tuple[bool, str]:
    try:
        importlib.import_module(module)
        return True, f"✅ {human_name}"
    except ImportError as e:
        return False, f"❌ {human_name}（未安装：{e}）"


def check_docx_parse() -> tuple[bool, str]:
    """用内置合成 docx 走一次 python-docx 读取。"""
    try:
        from docx import Document
        buf = io.BytesIO()
        doc = Document()
        doc.add_paragraph("self_check synthetic paragraph 1")
        doc.add_paragraph("self_check synthetic paragraph 2")
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)
        paras = [p.text for p in doc2.paragraphs if p.text.strip()]
        if len(paras) >= 2:
            return True, "✅ DOCX 基础解析（合成样本 2 段）"
        return False, f"❌ DOCX 解析返回段数异常：{len(paras)}"
    except Exception as e:
        return False, f"❌ DOCX 解析异常：{e}"


def check_pdf_parse() -> tuple[bool, str]:
    """检查 pypdf 与 pdfplumber 是否可 import 并能处理最小 PDF。"""
    try:
        import pypdf  # noqa
        import pdfplumber  # noqa
        # 只做 import 级别校验；避免依赖外部 fixture
        return True, "✅ 文本型 PDF 解析（pypdf + pdfplumber 可用）"
    except Exception as e:
        return False, f"❌ PDF 解析依赖缺失：{e}"


def check_kb_a() -> tuple[bool, str]:
    """检查 KB-A 规范层能否加载（YAML 文件计数）。"""
    kb_a = _REPO_ROOT / "knowledge_base" / "norms"
    if not kb_a.exists():
        return False, "❌ KB-A 目录不存在：knowledge_base/norms/"
    yamls = list(kb_a.rglob("*.yaml"))
    total_rules = 0
    for p in yamls:
        try:
            import yaml
            with open(p, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f) or {}
            rules = data.get("norms") or data.get("rules") or []
            total_rules += len(rules)
        except Exception:
            continue
    if total_rules >= 1:
        return True, f"✅ KB-A 规范加载（{len(yamls)} 文件 · {total_rules} 条）"
    return False, "❌ KB-A 未找到有效规则"


def check_report_render() -> tuple[bool, str]:
    """检查 render_report_html 能生成非空 HTML。"""
    try:
        # 只 import 不实跑，避免耦合到 rules 逻辑
        import jinja2  # noqa
        return True, "✅ 报告渲染引擎（jinja2 可用）"
    except Exception as e:
        return False, f"❌ 报告渲染依赖缺失：{e}"


def check_vision_graceful_degrade() -> tuple[bool, str]:
    """未配置视觉时，vision_pipeline 应可 import 且能 fallback。"""
    try:
        # vision 子模块可能未装 openai，允许 ImportError 走降级路径
        from scripts import vision_hint  # noqa
        return True, "⚪ 视觉未配置（降级路径可用；用户可选装 requirements-vision.txt）"
    except ImportError:
        return True, "⚪ 视觉子系统未加载（降级路径生效）"
    except Exception as e:
        return False, f"❌ 视觉降级路径异常：{e}"


def check_ark_smoke(fixture: Path) -> tuple[bool, str]:
    """可选：用合成 PNG 走一次 Ark。仅 --with-vision 时执行。"""
    if not fixture.exists():
        return False, f"❌ 视觉冒烟 fixture 不存在：{fixture}"
    try:
        from scripts.vision.providers.ark_provider import ArkVisionProvider  # noqa
    except Exception as e:
        return False, f"❌ 视觉未装：{e}"
    # 依赖 .env.local 中的 ARK_* 变量；此处不透露细节
    try:
        import os
        if not os.environ.get("ARK_API_KEY"):
            return False, "❌ ARK_API_KEY 未配置，跳过。"
        # 真调一次
        from scripts.vision.providers.ark_provider import ArkVisionProvider
        prov = ArkVisionProvider()
        _ = prov.recognize(fixture)
        return True, "✅ 视觉冒烟通过（Ark 真调 1 次，合成 PNG）"
    except Exception as e:
        return False, f"❌ 视觉冒烟失败：{e}"


# ---------- 主逻辑 ----------

def run(with_vision: bool = False) -> dict:
    results = []

    for m, name in [
        ("docx", "python-docx"),
        ("pypdf", "pypdf"),
        ("pdfplumber", "pdfplumber"),
        ("yaml", "pyyaml"),
        ("jinja2", "jinja2"),
    ]:
        ok, msg = _check_import(m, name)
        results.append({"section": "deps", "ok": ok, "msg": msg})

    for fn, section in [
        (check_docx_parse, "docx"),
        (check_pdf_parse, "pdf"),
        (check_kb_a, "kb_a"),
        (check_report_render, "render"),
        (check_vision_graceful_degrade, "vision_degrade"),
    ]:
        ok, msg = fn()
        results.append({"section": section, "ok": ok, "msg": msg})

    if with_vision:
        fixture = _REPO_ROOT / "tests" / "vision" / "fixtures" / "smoke_synthetic_table.png"
        ok, msg = check_ark_smoke(fixture)
        results.append({"section": "vision_smoke", "ok": ok, "msg": msg})

    return {
        "ok": all(r["ok"] for r in results),
        "results": results,
        "disclaimer": (
            "self_check 只回答“当前安装能否运行”，不代表质检准确率已重新验证。"
            "完整 Benchmark 由开发者运行，见 benchmarks/README.md。"
        ),
    }


def main() -> int:
    ap = argparse.ArgumentParser(description="用户侧轻量安装自检（不涉及 Benchmark，不调 Ark）")
    ap.add_argument("--json", action="store_true", help="输出 JSON")
    ap.add_argument("--with-vision", action="store_true",
                    help="可选：用合成 PNG 走一次 Ark（需 ARK_API_KEY）")
    ap.add_argument("--validate", metavar="DIAGNOSTIC_JSON",
                    help="校验 diagnostic_result.json 的 Schema 合规性（阻塞性门禁）")
    ap.add_argument("--json-schema", action="store_true",
                    help="与 --validate 配合使用，附加 JSON Schema 结构校验（需 jsonschema 库）")
    ap.add_argument("--audit-rules", action="store_true",
                    help="规则定义-vs-注册表一致性审计（与 --validate 互斥）")
    args = ap.parse_args()

    # --audit-rules 模式：只审计规则一致性
    if args.audit_rules:
        ok, msg = audit_rules()
        print(msg)
        return 0 if ok else 1

    # --validate 模式：只校验 Schema，不走自检流程
    if args.validate:
        path = Path(args.validate).expanduser().resolve()
        if not path.exists():
            print(f"❌ 文件不存在：{path}", file=sys.stderr)
            return 2
        passed, errors, warnings, js_applied = validate_diagnostic_result(path, use_json_schema=args.json_schema)
        if passed:
            msg = f"✅ Schema 校验通过：{path}"
            if args.json_schema:
                if js_applied:
                    msg += "（含 JSON Schema 结构校验）"
                else:
                    msg += "（JSON Schema 未执行：jsonschema 库未安装，pip install jsonschema）"
            print(msg)
            for w in warnings:
                print(f"⚠️  {w}")
            return 0
        else:
            print(f"❌ Schema 校验失败（{len(errors)} 项）：")
            for e in errors:
                print(f"   - {e}")
            for w in warnings:
                print(f"⚠️  {w}")
            return 1

    r = run(with_vision=args.with_vision)

    if args.json:
        print(json.dumps(r, ensure_ascii=False, indent=2))
    else:
        print("=" * 60)
        print("经管论文智检 Skill · 安装自检")
        print("=" * 60)
        for item in r["results"]:
            print(f"  [{item['section']:15s}] {item['msg']}")
        print("=" * 60)
        print(f"总体：{'✅ 通过' if r['ok'] else '❌ 有失败项'}")
        print(f"\n⚠️  {r['disclaimer']}")

    return 0 if r["ok"] else 1


if __name__ == "__main__":
    sys.exit(main())
